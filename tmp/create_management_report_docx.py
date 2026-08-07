from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"C:\Users\Glory\Documents\argleadstracker\output\ARG_Leads_Tracker_Management_Progress_and_Go_Live_Readiness_2026-08-04.docx")
NAVY = "06283D"
BLUE = "1F6AA5"
GREEN = "6D9F3D"
AMBER = "E6A933"
RED = "D94A48"
INK = "151515"
MUTED = "475569"
PALE_BLUE = "DCEAF6"
PALE_GREEN = "E6F0D8"
PALE_AMBER = "F8EDCE"
PALE_RED = "FAE4E2"


def set_cell_fill(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    tc_pr.append(shading)


def set_cell_border(cell, color="CBD5DF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def write_run(paragraph, text, size=10, bold=False, color=INK):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(11 if level == 1 else 7)
    paragraph.paragraph_format.space_after = Pt(4)
    size = 16 if level == 1 else 12
    write_run(paragraph, text, size=size, bold=True, color=NAVY)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        write_run(paragraph, bold_prefix, size=10, bold=True)
        write_run(paragraph, text[len(bold_prefix):], size=10, color=MUTED)
    else:
        write_run(paragraph, text, size=10, color=MUTED)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    write_run(paragraph, text, size=10, color=INK)
    return paragraph


def set_table_text(cell, text, color=INK, bold=False, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    write_run(paragraph, text, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_border(cell)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
normal.font.size = Pt(10)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
write_run(header, "AL RAS STEEL INTELLIGENCE | CONFIDENTIAL", size=8, bold=True, color=MUTED)

title_box = doc.add_table(rows=1, cols=1)
title_box.alignment = WD_TABLE_ALIGNMENT.CENTER
title_cell = title_box.cell(0, 0)
set_cell_fill(title_cell, NAVY)
set_cell_border(title_cell, NAVY)
title_cell.text = ""
title_p = title_cell.paragraphs[0]
title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
title_p.paragraph_format.space_after = Pt(2)
write_run(title_p, "ARG LEADS TRACKER CRM", size=20, bold=True, color="FFFFFF")
sub_p = title_cell.add_paragraph()
sub_p.paragraph_format.space_after = Pt(0)
write_run(sub_p, "Management Progress and Go-Live Readiness Report", size=11, bold=True, color="DCEAF6")
date_p = title_cell.add_paragraph()
date_p.paragraph_format.space_after = Pt(0)
write_run(date_p, "Prepared 4 August 2026 | Production-readiness status", size=9, color="FFFFFF")

add_heading(doc, "Executive Summary")
add_body(doc, "The CRM has completed its major frontend and security-hardening milestones. The Admin and Salesman experiences now share a consistent flat Bauhaus interface, core workflows remain intact, the production database was backed up, and the latest Supabase access-control migrations were applied and verified.")
add_body(doc, "Recommended release status: conditional go-live readiness. The product should only be released after the short operational checklist in this report is completed and signed off.")

add_heading(doc, "Current Readiness")
readiness = doc.add_table(rows=1, cols=3)
readiness.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = [("Area", NAVY), ("Status", NAVY), ("Management interpretation", NAVY)]
for index, (label, color) in enumerate(headers):
    set_cell_fill(readiness.cell(0, index), color)
    set_table_text(readiness.cell(0, index), label, color="FFFFFF", bold=True, size=9)
rows = [
    ("UI and workflow delivery", "Complete", "Admin and Salesman screens were modernized without intended business-logic changes.", PALE_GREEN),
    ("Database backup", "Complete", "Logical database export, manifests, checksums, and scheduled-backup verification are recorded.", PALE_GREEN),
    ("Database migrations and RLS", "Complete", "Schema reconciliation and policy repairs were applied; post-install checks passed.", PALE_GREEN),
    ("Automated validation", "Complete", "Security authorization tests and production build passed locally.", PALE_GREEN),
    ("Vercel production configuration", "Pending", "Required production secrets and scheduled-job configuration need confirmation.", PALE_AMBER),
    ("Final business UAT and rollback drill", "Pending", "Admin and Salesman acceptance testing must be documented before final sign-off.", PALE_AMBER),
]
for area, status, interpretation, fill in rows:
    cells = readiness.add_row().cells
    for cell in cells:
        set_cell_fill(cell, fill)
    set_table_text(cells[0], area, bold=True)
    set_table_text(cells[1], status, bold=True, color=GREEN if status == "Complete" else "8A5A00")
    set_table_text(cells[2], interpretation, color=MUTED)

add_heading(doc, "Delivered CRM Work")
for item in [
    "Unified Admin and Salesman visual system: flat Bauhaus palette, solid panels, accessible contrast, consistent headers, buttons, badges, tables, filters, and modals.",
    "Dashboard usability improvements: interactive KPI and pipeline-summary drill-downs, internal scrolling for long lists, responsive spacing, exports and date filters where required.",
    "Lead, Pipeline, Activity, Tasks, Salesmen, Accounts, Quotes, and Weekly Sales Report readability and responsive-layout refinements.",
    "Role-preserving workflow improvements, including actionable summary dialogs, lead-detail readability fixes, and activity/task interaction enhancements.",
    "AI Sales Assistant presentation and workflow integration, retaining existing role-based CRM behavior.",
]:
    add_bullet(doc, item)

add_heading(doc, "Security and Recovery Evidence")
for item in [
    "Production Supabase logical backup verified at C:\\Users\\Glory\\Documents\\local-release-files\\supabase-backup-20260803-132551.",
    "Backup contents: roles.sql, schema.sql, data.sql, BACKUP-MANIFEST.txt, SHA256SUMS.txt, plus a Storage inventory. Storage buckets were inspected; lead-files and pmr-voice-notes had zero objects at capture time.",
    "Supabase migrations applied: production schema reconciliation, legacy notifications-schema reconciliation, and legacy RLS policy replacement.",
    "Post-install database checks confirmed the required rate-limit objects, lead access policies, Storage access policies, user-scoped RLS helpers, and weekly-report lifecycle controls.",
    "Eight authorization-boundary tests passed. They cover signed-in-user REST access, missing-auth rejection, constrained service-role access, cron authorization and idempotency, JWT-bound storage upload/signing, and private CRM media controls.",
    "npm run build completed successfully.",
]:
    add_bullet(doc, item)

add_heading(doc, "Mandatory Steps Before Final Go-Live")
pending = doc.add_table(rows=1, cols=4)
pending.alignment = WD_TABLE_ALIGNMENT.CENTER
for index, label in enumerate(["Priority", "Action", "Owner", "Completion evidence"]):
    set_cell_fill(pending.cell(0, index), NAVY)
    set_table_text(pending.cell(0, index), label, color="FFFFFF", bold=True, size=9)
items = [
    ("P0", "Configure and verify Vercel production secrets, including RATE_LIMIT_HASH_SECRET and CRON_SECRET. Confirm that scheduled jobs no longer return 401/503.", "Platform owner", "Vercel environment screenshot/redacted verification and successful protected-job response.", PALE_RED),
    ("P0", "Run formal UAT with one Admin and one Salesman account: login, role navigation, dashboard counts, filters, lead detail, task/report draft and submit, voice-note paths, exports, and AI assistant access.", "Business owner + QA", "Signed UAT checklist with issues resolved or accepted.", PALE_RED),
    ("P0", "Perform a restore rehearsal into a non-production Supabase project from the current logical export. Do not restore into production.", "Database owner", "Restore log and validation of representative tables and role policies.", PALE_RED),
    ("P1", "Enable release monitoring and document the owner/escalation path for Vercel errors, Supabase errors, cron failures, and failed user workflows.", "Platform owner", "Alert settings and named on-call/escalation contacts.", PALE_AMBER),
    ("P1", "Confirm Vercel deployment health through live TLS route checks after deployment, not only the dashboard Ready state.", "Release owner", "HTTP smoke-test record for the production alias and critical authenticated flows.", PALE_AMBER),
    ("P1", "Schedule periodic database and Storage backup verification. Storage backups are separate from database logical backups.", "Database owner", "Documented schedule, retention owner, and test-recovery cadence.", PALE_AMBER),
]
for priority, action, owner, evidence, fill in items:
    cells = pending.add_row().cells
    for cell in cells:
        set_cell_fill(cell, fill)
    set_table_text(cells[0], priority, bold=True, color=RED if priority == "P0" else "8A5A00")
    set_table_text(cells[1], action)
    set_table_text(cells[2], owner, bold=True)
    set_table_text(cells[3], evidence, color=MUTED)

add_heading(doc, "Release Recommendation")
add_body(doc, "Do not treat the current state as an unconditional final launch until the P0 tasks are completed. The application is suitable for a controlled production deployment after Vercel secrets, live protected-job validation, business UAT, and a non-production restore rehearsal are complete.")

note = doc.add_table(rows=1, cols=1)
note_cell = note.cell(0, 0)
set_cell_fill(note_cell, PALE_BLUE)
set_cell_border(note_cell, BLUE)
note_cell.text = ""
note_p = note_cell.paragraphs[0]
write_run(note_p, "Security note: ", size=9, bold=True, color=NAVY)
write_run(note_p, "This report intentionally excludes passwords, API keys, database URLs, tokens, and personal secrets.", size=9, color=INK)

footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
write_run(footer, "ARG Leads Tracker CRM | Management Report | 4 August 2026", size=8, color=MUTED)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
